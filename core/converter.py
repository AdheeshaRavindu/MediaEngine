import shutil
import subprocess
import tempfile
from pathlib import Path

from core.utils import build_output_folder, build_output_path, detect_file_type, get_tool_path, is_python_module_available, is_tool_available
from engines import ffmpeg_wrapper, imagemagick_wrapper


IMAGE_FORMATS = {"png", "jpg", "jpeg", "webp", "bmp", "tiff", "tif", "pdf"}
VIDEO_FORMATS = {"mp4", "mkv", "avi", "mov", "webm"}
AUDIO_FORMATS = {"mp3", "wav", "aac", "flac", "ogg", "m4a"}
DOCUMENT_FORMATS = {"doc", "docx", "docm", "dot", "dotx", "dotm", "rtf", "odt"}
DOC_IMAGE_FORMATS = {"png", "jpg", "jpeg", "webp", "bmp", "tiff", "tif"}
DOCX2PDF_COMPATIBLE_FORMATS = {"doc", "docx", "docm", "dot", "dotx", "dotm", "rtf"}
DOCX_BASIC_PDF_COMPATIBLE_FORMATS = {"docx"}


def _load_fitz():
    try:
        import fitz
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "PyMuPDF is not installed for the Python interpreter running this app. "
            "Run: py -m pip install -r requirements.txt"
        ) from exc
    return fitz


def _normalize_ext(fmt: str | None) -> str | None:
    if not fmt:
        return None
    value = fmt.strip().lower()
    if value == "auto":
        return None
    if value.startswith("."):
        value = value[1:]
    return value or None


def _office_executable() -> str:
    for tool_name in ("soffice", "libreoffice"):
        if is_tool_available(tool_name):
            return get_tool_path(tool_name)
    raise RuntimeError("LibreOffice is required for Word document conversions. Install soffice or libreoffice.")


def _convert_with_office(input_path: str, output_dir: str, target_ext: str) -> str:
    source = Path(input_path)
    output_path = build_output_path(input_path, output_dir, "_converted", f".{target_ext}")

    if source.suffix.lower().lstrip(".") == target_ext:
        shutil.copy2(input_path, output_path)
        return output_path

    with tempfile.TemporaryDirectory() as temp_dir:
        command = [
            _office_executable(),
            "--headless",
            "--nologo",
            "--convert-to",
            target_ext,
            "--outdir",
            temp_dir,
            input_path,
        ]
        subprocess.run(command, check=True, capture_output=True, text=True)

        converted = Path(temp_dir) / f"{source.stem}.{target_ext}"
        if not converted.exists():
            candidates = list(Path(temp_dir).glob(f"*.{target_ext}"))
            if not candidates:
                raise RuntimeError(f"LibreOffice did not produce a .{target_ext} file")
            converted = candidates[0]

        shutil.move(str(converted), output_path)

    return output_path


def _convert_with_docx2pdf(input_path: str, output_dir: str) -> str:
    try:
        from docx2pdf import convert as docx2pdf_convert
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "LibreOffice is not available and docx2pdf is not installed. "
            "Install LibreOffice, or run: py -m pip install docx2pdf"
        ) from exc

    source_ext = Path(input_path).suffix.lower().lstrip(".")
    if source_ext not in DOCX2PDF_COMPATIBLE_FORMATS:
        raise RuntimeError(
            f"This .{source_ext} file needs LibreOffice for PDF export. "
            "docx2pdf fallback supports DOC/DOCX/DOCM/DOT/DOTX/DOTM/RTF with Microsoft Word."
        )

    output_path = build_output_path(input_path, output_dir, "_converted", ".pdf")
    try:
        docx2pdf_convert(input_path, output_path)
    except Exception as exc:
        raise RuntimeError(
            "DOCX to PDF fallback failed. Ensure Microsoft Word is installed, or install LibreOffice."
        ) from exc

    if not Path(output_path).exists():
        raise RuntimeError("DOCX to PDF fallback did not produce an output file")

    return output_path


def _convert_docx_to_pdf_basic(input_path: str, output_dir: str) -> str:
    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "python-docx is required for basic DOCX to PDF fallback. "
            "Run: py -m pip install python-docx"
        ) from exc

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "reportlab is required for basic DOCX to PDF fallback. "
            "Run: py -m pip install reportlab"
        ) from exc

    doc = Document(input_path)
    output_path = build_output_path(input_path, output_dir, "_converted", ".pdf")

    pdf = canvas.Canvas(output_path, pagesize=A4)
    page_width, page_height = A4
    margin = 48
    line_height = 14
    max_chars = 105
    y = page_height - margin

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            y -= line_height
            if y <= margin:
                pdf.showPage()
                y = page_height - margin
            continue

        remaining = text
        while remaining:
            line = remaining[:max_chars]
            remaining = remaining[max_chars:]
            pdf.drawString(margin, y, line)
            y -= line_height
            if y <= margin:
                pdf.showPage()
                y = page_height - margin

    pdf.save()

    if not Path(output_path).exists():
        raise RuntimeError("Basic DOCX to PDF fallback did not produce an output file")

    return output_path


def _pdf_to_images(input_pdf: str, output_dir: str, image_ext: str) -> str:
    output_folder = build_output_folder(input_pdf, output_dir, f"_pages_{image_ext}")
    Path(output_folder).mkdir(parents=True, exist_ok=True)

    fitz = _load_fitz()
    document = fitz.open(input_pdf)
    try:
        zoom = fitz.Matrix(2.0, 2.0)
        for index, page in enumerate(document, start=1):
            pixmap = page.get_pixmap(matrix=zoom, alpha=False)
            page_path = Path(output_folder) / f"page-{index:03d}.{image_ext}"
            pixmap.save(str(page_path))
    finally:
        document.close()

    if not list(Path(output_folder).glob(f"*.{image_ext}")):
        raise RuntimeError("PyMuPDF did not produce any page images")

    return output_folder


def _image_files_to_docx(image_paths: list[Path], output_path: str) -> str:
    if not image_paths:
        raise ValueError("No images were found for DOCX conversion")

    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "python-docx is not installed for the Python interpreter running this app. "
            "Run: py -m pip install -r requirements.txt"
        ) from exc

    document = Document()
    section = document.sections[0]
    max_width = section.page_width - section.left_margin - section.right_margin

    for index, image_path in enumerate(image_paths):
        if index > 0:
            document.add_page_break()
        document.add_picture(str(image_path), width=max_width)

    document.save(output_path)
    return output_path


def _pdf_or_image_to_docx(input_path: str, output_dir: str) -> str:
    output_path = build_output_path(input_path, output_dir, "_converted", ".docx")
    source_type = detect_file_type(input_path)

    if source_type == "image":
        return _image_files_to_docx([Path(input_path)], output_path)

    if source_type == "pdf":
        with tempfile.TemporaryDirectory() as temp_dir:
            pages_folder = _pdf_to_images(input_path, temp_dir, "png")
            image_paths = sorted(Path(pages_folder).glob("page-*.png"))
            return _image_files_to_docx(image_paths, output_path)

    raise ValueError("DOCX export supports image and PDF inputs only")


def _document_to_image_folder(input_path: str, output_dir: str, image_ext: str) -> str:
    with tempfile.TemporaryDirectory() as temp_dir:
        pdf_path = _convert_with_office(input_path, temp_dir, "pdf")
        return _pdf_to_images(pdf_path, output_dir, image_ext)


def run_convert(input_path: str, output_dir: str, output_format: str | None = None) -> str:
    file_type = detect_file_type(input_path)
    target_ext = _normalize_ext(output_format)

    if file_type in {"video", "audio"}:
        if target_ext is None:
            target_ext = "mp4" if file_type == "video" else "mp3"
        if target_ext not in VIDEO_FORMATS and target_ext not in AUDIO_FORMATS:
            raise ValueError(f"Unsupported media output format: {target_ext}")
        output_path = build_output_path(input_path, output_dir, "_converted", f".{target_ext}")
        ffmpeg_wrapper.transcode_media(input_path, output_path)
        return output_path

    if file_type == "document":
        if target_ext is None:
            target_ext = "pdf"
        if target_ext == "pdf":
            source_ext = Path(input_path).suffix.lower().lstrip(".")
            if is_tool_available("soffice") or is_tool_available("libreoffice"):
                return _convert_with_office(input_path, output_dir, target_ext)
            if source_ext in DOCX2PDF_COMPATIBLE_FORMATS and is_python_module_available("docx2pdf"):
                try:
                    return _convert_with_docx2pdf(input_path, output_dir)
                except RuntimeError:
                    # If Word automation fails, fall back to basic text export for .docx.
                    if source_ext in DOCX_BASIC_PDF_COMPATIBLE_FORMATS and is_python_module_available("docx") and is_python_module_available("reportlab"):
                        return _convert_docx_to_pdf_basic(input_path, output_dir)
                    raise
            if source_ext in DOCX_BASIC_PDF_COMPATIBLE_FORMATS and is_python_module_available("docx") and is_python_module_available("reportlab"):
                return _convert_docx_to_pdf_basic(input_path, output_dir)
            raise RuntimeError(
                "DOCX to PDF requires one of: LibreOffice, docx2pdf (with Microsoft Word), or basic fallback (python-docx + reportlab for .docx)."
            )
        if target_ext in DOCUMENT_FORMATS:
            return _convert_with_office(input_path, output_dir, target_ext)
        if target_ext in DOC_IMAGE_FORMATS:
            return _document_to_image_folder(input_path, output_dir, target_ext)
        raise ValueError(f"Unsupported Word/document output format: {target_ext}")

    if file_type in {"image", "pdf"}:
        if target_ext is None:
            target_ext = "jpg" if file_type == "image" else "pdf"
        if target_ext == "docx":
            return _pdf_or_image_to_docx(input_path, output_dir)
        if file_type == "pdf" and target_ext in DOC_IMAGE_FORMATS:
            return _pdf_to_images(input_path, output_dir, target_ext)
        if target_ext not in IMAGE_FORMATS:
            raise ValueError(f"Unsupported image/pdf output format: {target_ext}")
        output_path = build_output_path(input_path, output_dir, "_converted", f".{target_ext}")
        imagemagick_wrapper.convert_image(input_path, output_path)
        return output_path

    raise ValueError("Unsupported file type for convert")
